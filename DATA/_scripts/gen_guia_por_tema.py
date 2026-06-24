# -*- coding: utf-8 -*-
# Genera la GUÍA POR TEMA (de dónde repasar + ángulos más preguntados + trampa) desde el resultado
# del workflow encaps-guia-por-tema → MD (DATA/ENCAPS) + Word (ACTUAL). python3 gen_guia_por_tema.py
import json, os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUTFILE = r'C:\Users\JOSEPH~1\AppData\Local\Temp\claude\D--joseph-md-app\2b1d4275-ceb5-4f8d-9ab0-d4dbbd76c52a\tasks\wi3k2sch0.output'
d = json.load(open(OUTFILE, encoding='utf-8'))

def find_areas(o):
    if isinstance(o, list) and o and isinstance(o[0], dict) and 'area' in o[0] and 'temas' in o[0]: return o
    if isinstance(o, dict):
        for v in o.values():
            r = find_areas(v)
            if r: return r
    if isinstance(o, list):
        for v in o:
            r = find_areas(v)
            if r: return r
    return None
areas = find_areas(d)

# ---- 1) Markdown (para el generador de preguntas + futuros chats) ----
md = ['# 📖 GUÍA POR TEMA ENCAPS 2026-II — de dónde repasar + qué se pregunta más',
      '',
      '> Generado del análisis de los **6 exámenes oficiales reales** (exams_txt) + fuentes (encapsFuentes / encapsVideosPorTema) por 5 agentes (1 por área).',
      '> Para CADA tema: **dónde estudiarlo** (fuentes concretas), los **ángulos más preguntados** (citando el examen) y la **trampa clásica** del distractor.',
      '> Úsalo junto con `GENERADOR_PREGUNTAS_ENCAPS.md` (el generador prioriza estos ángulos).', '']
for a in areas:
    md.append(f'## {a["area"]}')
    md.append('')
    for t in a['temas']:
        md.append(f'### {t["codigo"]} · {t["subtema"]}')
        md.append('**📚 De dónde repasar:**')
        for s in t.get('deDondeEstudiar', []): md.append(f'- {s}')
        md.append('')
        md.append('**🎯 Ángulos más preguntados (exámenes reales):**')
        for s in t.get('angulos', []): md.append(f'- {s}')
        md.append('')
        md.append(f'**⚠️ Trampa clásica:** {t.get("trampa","")}')
        md.append('')
out_md = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), 'ENCAPS', 'GUIA_POR_TEMA_2026-2.md')
open(out_md, 'w', encoding='utf-8').write('\n'.join(md))
print('MD:', out_md, '·', sum(len(a['temas']) for a in areas), 'temas')

# ---- 2) Word élite (para estudiar) ----
NAVY = RGBColor(0x1A,0x2B,0x4A); GOLD = RGBColor(0xB8,0x86,0x0B); RED = RGBColor(0xC0,0x39,0x2B); GREEN = RGBColor(0x1E,0x7E,0x34); GREY = RGBColor(0x55,0x55,0x55)
doc = Document(); st = doc.styles['Normal']; st.font.name = 'Calibri'; st.font.size = Pt(10.5)
def H(txt, sz=14, c=NAVY, af=6, bf=10):
    p = doc.add_paragraph(); r = p.add_run(txt); r.bold = True; r.font.size = Pt(sz); r.font.color.rgb = c
    p.paragraph_format.space_after = Pt(af); p.paragraph_format.space_before = Pt(bf)
def P(txt, bold=False, c=None, sz=10.5, bullet=False):
    p = doc.add_paragraph(style='List Bullet' if bullet else None); r = p.add_run(txt); r.bold = bold; r.font.size = Pt(sz)
    if c: r.font.color.rgb = c
    p.paragraph_format.space_after = Pt(2)
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('GUÍA POR TEMA ENCAPS 2026-II'); r.bold = True; r.font.size = Pt(20); r.font.color.rgb = NAVY
s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = s.add_run('De dónde repasar · ángulos más preguntados (6 exámenes reales) · trampa clásica'); r.bold = True; r.font.size = Pt(11); r.font.color.rgb = GOLD
for a in areas:
    H(a['area'], 15, GOLD, bf=14)
    for t in a['temas']:
        H(f'{t["codigo"]} · {t["subtema"]}', 12, NAVY, af=3, bf=8)
        P('📚 De dónde repasar:', bold=True)
        for x in t.get('deDondeEstudiar', []): P(x, bullet=True, sz=9.5)
        P('🎯 Ángulos más preguntados:', bold=True, c=GREEN)
        for x in t.get('angulos', []): P(x, bullet=True, sz=9.5)
        P(f'⚠️ Trampa clásica: {t.get("trampa","")}', c=RED, sz=9.5)
out_dir = r'D:\agente_estudio\ENCAPS\ENCAPS\FORMAS DE ESTUDIO\ACTUAL'
os.makedirs(out_dir, exist_ok=True)
out_docx = os.path.join(out_dir, 'Guia_POR_TEMA_ENCAPS_2026-2_v10.docx')
doc.save(out_docx)
print('Word:', out_docx)
