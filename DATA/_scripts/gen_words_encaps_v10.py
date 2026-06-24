# -*- coding: utf-8 -*-
# Genera 3 Words ENCAPS v10 (DATA-DRIVEN) en D:\agente_estudio\ENCAPS\...\ACTUAL\:
#  1) RENTABILIDAD por pilar + por tema (pronóstico 2026-2)
#  2) MAESTRO global 25-jun→20-ago (Fase1 ≤17d TODO el temario front-load Pareto · Fase2 vueltas/preguntas · Fase3 sims)
#  3) SEMANAL detallado 25-28 jun
# python3 gen_words_encaps_v10.py
import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = r'D:\agente_estudio\ENCAPS\ENCAPS\FORMAS DE ESTUDIO\ACTUAL'
os.makedirs(OUT, exist_ok=True)
NAVY=RGBColor(0x1A,0x2B,0x4A); GOLD=RGBColor(0xB8,0x86,0x0B); RED=RGBColor(0xC0,0x39,0x2B)
GREEN=RGBColor(0x1E,0x7E,0x34); GREY=RGBColor(0x55,0x55,0x55)

# ---- TABLA MAESTRA: 40 temas (codigo, subtema, area, prioridad, rank, pct, vueltas) data-driven ----
A_SP='Salud Pública'; A_CI='Cuidado Integral'; A_ET='Ética/Intercultural'; A_IN='Investigación'; A_GE='Gestión'
V={'CRÍTICA':6,'ALTA':5,'MEDIA':4,'BAJA':3}
TEMAS=[
 ('I-3','Vigilancia epidemiológica / brotes',A_SP,'CRÍTICA',1,10.5),
 ('V-2','Planeamiento institucional PEI/POI/FODA',A_GE,'CRÍTICA',2,11.0),
 ('I-5+I-6','Determinantes sociales + Bioestadística',A_SP,'CRÍTICA',3,7.0),
 ('II-3','Esquema nacional de vacunación / ESAVI',A_CI,'CRÍTICA',4,5.5),
 ('III-5','Salud intercultural / medicina tradicional',A_ET,'ALTA',5,4.5),
 ('I-4','Definiciones de caso (dengue/TB/malaria/metaxénicas)',A_SP,'ALTA',6,4.0),
 ('V-3','Niveles de atención y referencia',A_GE,'ALTA',7,4.0),
 ('II-11','ITS / VIH / sífilis',A_CI,'ALTA',8,3.8),
 ('I-1','Promoción de la salud',A_SP,'ALTA',9,3.5),
 ('II-1','Salud materna / prenatal / parto',A_CI,'ALTA',10,3.5),
 ('II-8','HTA / Diabetes / ECNT / HEARTS',A_CI,'ALTA',11,3.0),
 ('V-1','Categorización de EESS',A_GE,'ALTA',12,3.0),
 ('III-9','Derechos del paciente',A_ET,'ALTA',13,3.0),
 ('IV-1+IV-2','Tipos/diseños de estudio + Validez',A_IN,'ALTA',14,3.0),
 ('I-2','FESP renovadas',A_SP,'MEDIA',15,2.5),
 ('III-2','CMP / código deontológico / acto médico',A_ET,'MEDIA',16,2.5),
 ('III-8','Ética de la función pública (Ley 27815)',A_ET,'MEDIA',17,2.5),
 ('II-6','Tuberculosis',A_CI,'MEDIA',18,2.3),
 ('II-9','Salud mental',A_CI,'MEDIA',19,2.2),
 ('II-7','Adulto mayor / VACAM',A_CI,'MEDIA',20,2.0),
 ('II-5','Modelo de Cuidado Integral / paquetes',A_CI,'MEDIA',21,2.0),
 ('II-10','Cáncer / oncología',A_CI,'MEDIA',22,2.0),
 ('II-4','Anemia y nutrición',A_CI,'MEDIA',23,1.8),
 ('III-1','Bioética / principios',A_ET,'MEDIA',24,1.5),
 ('III-4+III-7','Violencia género/familiar + Aborto/dilemas',A_ET,'MEDIA',25,1.5),
 ('I-11+I-12','Salud familiar + comunitaria',A_SP,'MEDIA',26,1.3),
 ('V-6','Telesalud',A_GE,'MEDIA',27,1.3),
 ('III-6+III-10','Política intercultural / adecuación',A_ET,'BAJA',28,1.2),
 ('I-7','PNAIA niñez / adolescencia',A_SP,'BAJA',29,1.0),
 ('II-2','CRED',A_CI,'BAJA',30,1.0),
 ('I-10','APS específica / Alma-Ata',A_SP,'BAJA',31,0.9),
 ('III-3','Consentimiento informado',A_ET,'BAJA',32,0.9),
 ('V-7+V-10','NTS Referencia + SIS/AUS/aseguramiento',A_GE,'BAJA',33,0.9),
 ('IV-4','OR/RR/medidas de asociación',A_IN,'BAJA',35,0.6),
 ('IV-3+IV-5','Tamizaje + Pruebas diagnósticas',A_IN,'BAJA',36,0.5),
 ('II-12','Salud bucal',A_CI,'BAJA',37,0.5),
 ('IV-6+IV-7','Indicadores / sistema de vigilancia',A_IN,'BAJA',39,0.3),
 ('I-8','PNDH discapacidad',A_SP,'BAJA',40,0.5),
 ('I-9','ENSF familias',A_SP,'BAJA',41,0.5),
 ('II-13','Salud ocular y oído',A_CI,'BAJA',42,0.4),
]
PILAR_PCT={A_SP:29,A_CI:28,A_GE:21,A_ET:16,A_IN:6}

# ---- FASE-1 front-load: todo el temario en los primeros ~18 días hábiles, Pareto (CRÍTICA 1/día, BAJA agrupada) ----
# (codigos por día-hábil; el día-calendario lo maneja la app/Calendar)
FASE1=[
 ('D1 · Jue 25-jun',['I-3']),('D2 · Vie 26-jun',['V-2']),
 ('D4 · Lun 29-jun',['I-5+I-6']),('D5 · Mar 30-jun',['II-3']),('D6 · Mié 1-jul',['III-5','II-6']),
 ('D7 · Jue 2-jul',['I-4','II-9']),('D8 · Vie 3-jul',['V-3','II-7']),
 ('D10 · Lun 6-jul',['II-11','II-5']),('D11 · Mar 7-jul',['I-1','II-10']),
 ('D12 · Mié 8-jul',['II-1','II-4']),('D13 · Jue 9-jul',['II-8','III-1']),
 ('D14 · Vie 10-jul',['V-1','III-4+III-7']),('D16 · Lun 13-jul',['III-9','I-11+I-12','V-6']),
 ('D17 · Mar 14-jul',['IV-1+IV-2','III-6+III-10','I-7']),
 ('D18 · Mié 15-jul',['I-2','II-2','I-10','III-3']),
 ('D19 · Jue 16-jul',['III-2','V-7+V-10','IV-4','IV-3+IV-5']),
 ('D20 · Vie 17-jul',['III-8','II-12','IV-6+IV-7','I-8','I-9','II-13']),
]
SUB={t[0]:t[1] for t in TEMAS}; AR={t[0]:t[2] for t in TEMAS}; PR={t[0]:t[3] for t in TEMAS}

def newdoc():
    d=Document(); s=d.styles['Normal']; s.font.name='Calibri'; s.font.size=Pt(10.5); return d
def H(d,txt,sz=14,c=NAVY,af=6,bf=10):
    p=d.add_paragraph(); r=p.add_run(txt); r.bold=True; r.font.size=Pt(sz); r.font.color.rgb=c
    p.paragraph_format.space_after=Pt(af); p.paragraph_format.space_before=Pt(bf); return p
def P(d,txt,bold=False,c=None,sz=10.5,bullet=False):
    p=d.add_paragraph(style='List Bullet' if bullet else None); r=p.add_run(txt); r.bold=bold; r.font.size=Pt(sz)
    if c: r.font.color.rgb=c
    p.paragraph_format.space_after=Pt(2); return p
def cover(d,title,sub):
    t=d.add_paragraph(); t.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=t.add_run(title); r.bold=True; r.font.size=Pt(20); r.font.color.rgb=NAVY
    s=d.add_paragraph(); s.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=s.add_run(sub); r.bold=True; r.font.size=Pt(12); r.font.color.rgb=GOLD
    for ln in ['Joseph Max Soto Tocas · Médico, Huancayo','🎯 META ≥17/20 · 🏥 Plaza Huachac, Junín · 📅 EXAMEN jue 20-ago-2026 (FIJO)',
               'DATA-DRIVEN: 6 exámenes oficiales reales (2024-II→2026-I, ~600 preguntas) + backtest walk-forward + ciencia del aprendizaje (Oakley·Newport·Make It Stick·Ultralearning·Ebbinghaus).']:
        p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; rr=p.add_run(ln); rr.font.size=Pt(9.5); rr.font.color.rgb=GREY

# =================== DOC 1 · RENTABILIDAD ===================
d=newdoc()
cover(d,'RENTABILIDAD ENCAPS 2026-II','Pronóstico data-driven por pilar y por tema · v10')
H(d,'1 · Rentabilidad por PILAR (proyección 2026-II)',13)
tb=d.add_table(rows=1,cols=3); tb.style='Light Grid Accent 1'
hc=tb.rows[0].cells; hc[0].text='Pilar (área)'; hc[1].text='% del examen'; hc[2].text='Estrategia'
for a,strat in [(A_SP,'MÁXIMO foco — vigilancia y bioestadística'),(A_CI,'ALTO — vacunas, ITS, materna, crónicas'),
                (A_GE,'ALTO — PEI/POI, categorización, referencia'),(A_ET,'MEDIO — intercultural, deberes/derechos'),
                (A_IN,'MÍNIMO — solo diseños (IV-1); resto en caída')]:
    c=tb.add_row().cells; c[0].text=a; c[1].text=f'{PILAR_PCT[a]}%'; c[2].text=strat
P(d,'')
P(d,'CI + SP ≈ 57% (NO 70%). Investigación cayó de 16% a 3-6%. Confianza: ranking por pilar fiable; % por tema = bandas ±2-4pp (n=4 exámenes).',c=GREY,sz=9.5)

H(d,'2 · Rentabilidad por TEMA dentro de cada pilar (rank · % · vueltas)',13)
for area in [A_SP,A_CI,A_GE,A_ET,A_IN]:
    H(d,f'▸ {area} ({PILAR_PCT[area]}% del examen)',11,GOLD,af=3,bf=8)
    tb=d.add_table(rows=1,cols=5); tb.style='Light List Accent 1'
    hc=tb.rows[0].cells
    for i,h in enumerate(['Cód','Tema','Rank','% est.','Vueltas']): hc[i].text=h
    for cod,subt,ar,pr,rk,pct in sorted([t for t in TEMAS if t[2]==area],key=lambda x:x[4]):
        c=tb.add_row().cells; c[0].text=cod; c[1].text=subt; c[2].text=f'#{rk}'; c[3].text=f'{pct}%'; c[4].text=f'{V[pr]} ({pr})'
    P(d,'')

H(d,'3 · Cómo se construyó (backtest walk-forward honesto)',12)
for x in ['2024-II → predijo 2025-I (error 3.2pp) → corrigió.','+2025-I → predijo 2025-II (5.6pp).','+2025-II → predijo 2026-I (4.1pp).',
          'TODO + aprendizajes → pronóstico 2026-II.','Aprendizajes: IV no se sostiene; I-3 (vigilancia) tendencia alcista fuerte 2.8→6.5→9→14%; picos por-tema (II-3, V-3) emergen sin aviso.']:
    P(d,x,bullet=True,sz=9.5)
P(d,'Data cruda: D:\\joseph-md-app\\DATA\\ENCAPS\\ (exams_txt/, _pronostico_2026_2.json, README_RENTABILIDAD_2026-2.md). Hallazgo: clúster Gestión-medicamentos (SISMED/PNUME/DIGEMID) ~6-12% sin código propio.',c=GREY,sz=9)
d.save(os.path.join(OUT,'Plan_RENTABILIDAD_ENCAPS_2026-2_v10.docx'))

# =================== DOC 2 · MAESTRO GLOBAL ===================
d=newdoc()
cover(d,'PLAN MAESTRO ENCAPS 2026-II','25-jun → 20-ago · 49 días hábiles · v10 DATA-DRIVEN (Pareto + repetición espaciada)')
H(d,'0 · Filosofía del plan (por qué v10)',13)
for x in ['REGLA DE ORO: TODO el temario (40 unidades-tema) se ve en los PRIMEROS ~18 días hábiles (Fase 1). Después NO se estudia temario nuevo — solo vueltas, mapas conceptuales, banco de preguntas y simulacros.',
          'Pareto 80/20: 80% del tiempo a las áreas/temas más rentables (SP+CI+Gestión = 78% del examen); los CRÍTICA tienen su propio día de deep-work; los BAJA se agrupan varios por día.',
          'Las plataformas (QX/Theomed) abren su 2ª fase (mapas conceptuales/banco) el 1-jul; nosotros adelantamos: temario terminado antes para empalmar con esa fase de preguntas.',
          'Repetición espaciada diferenciada + intervalos comprimidos para que toda vuelta caiga antes del 20-ago.']:
    P(d,x,bullet=True)

H(d,'1 · Aritmética de los 50 días',12)
tb=d.add_table(rows=1,cols=3); tb.style='Light Grid Accent 1'
hc=tb.rows[0].cells; hc[0].text='Fase'; hc[1].text='Días'; hc[2].text='Objetivo'
for f,dd,ob in [('FASE 1 · COBERTURA','25-jun → 17-jul (~17 días háb.)','1ª vuelta de TODO el temario, front-load Pareto. Cada tema cierra con banco (retrieval).'),
                ('FASE 2 · VUELTAS + PREGUNTAS','18-jul → 5-ago (~16 días háb.)','Repaso espaciado (2ª-6ª vuelta) + interleaving + mapas conceptuales + banco mezclado. 65-70% a alto yield. Sábados=simulacros.'),
                ('FASE 3 · RECTA FINAL','6-ago → 19-ago (~10 días háb.)','SOLO simulacros tamaño-examen + preguntas + estrategia. Cero temario nuevo. Sueño la víspera.'),
                ('EXAMEN','jue 20-ago','—')]:
    c=tb.add_row().cells; c[0].text=f; c[1].text=dd; c[2].text=ob
P(d,'')
P(d,'Domingos LIBRES (8) = consolidación por sueño/modo difuso. Sábados = simulacros con revisión brutal.',c=GREY,sz=9.5)

H(d,'2 · FASE 1 — Cobertura total del temario (día por día, front-load por rentabilidad)',13)
P(d,'CRÍTICA: 1 tema/día (deep work completo). ALTA: ~2/día. MEDIA/BAJA: agrupados. TODO el temario cubierto al día 18.',c=GREY,sz=9.5)
tb=d.add_table(rows=1,cols=3); tb.style='Light List Accent 1'
hc=tb.rows[0].cells; hc[0].text='Día'; hc[1].text='Temas (código)'; hc[2].text='Detalle'
for dia,cods in FASE1:
    det=' · '.join(f'{c} {SUB.get(c,"")} [{PR.get(c,"")[:4]}]' for c in cods)
    c=tb.add_row().cells; c[0].text=dia.split(' · ')[0]; c[1].text=', '.join(cods); c[2].text=det[:120]
P(d,'')
P(d,'Días 3,9,15,21 (sábados) = simulacros · Dom libres. Tras el día 20 (vie 17-jul) → NO más temario nuevo.',bold=True,c=RED,sz=10)

H(d,'3 · Sistema de vueltas (repetición espaciada diferenciada)',12)
for x in ['CRÍTICA = 6 vueltas, intervalos [1,3,7,28,50] · ALTA = 5 [1,7,28,50] · MEDIA = 4 [3,28,50] · BAJA = 3 [7,50].',
          'Como TODO se siembra en días 1-18, las cadenas de repaso de los CRÍTICA (sembrados días 1-5) completan sus 6 vueltas con espaciado real antes del 20-ago.',
          'Cada vuelta = RETRIEVAL (banco de preguntas de memoria), no relectura. Cada fallo resetea el subtema a 1-3 días.',
          'Motor: intervalosComprimidos() comprime las vueltas tardías para que ninguna caiga después del examen.']:
    P(d,x,bullet=True,sz=10)

H(d,'4 · Horario diario base (8 bloques · sincronizado con Google Calendar · los horarios no se mueven, solo el tema)',12)
for hr,bl in [('04:15-04:45','🟢 Anclaje Post-Sueño (mapa del tema previo + Anki)'),('04:45-05:45','🟡 20 Preguntas warm-up adaptativo'),
              ('07:15-08:15','🍅 Repaso Espaciado Multi-Temporal (D-1/D-3/D-7)'),('08:15-09:00','🟡 PRE-TEST 10Q ciegas + free recall'),
              ('09:00-11:00','🔴 NÚCLEO DEEP PRIME (videos QX + compendio + NTS + APEX) — solo Fase 1'),('11:00-12:00','🟡 30 Preguntas Consolidación'),
              ('17:15-18:00','🟢 Anclaje Vespertino'),('18:00-18:45','🔥 Evaluación Diaria Acumulativa + POSTEST Theomed')]:
    p=d.add_paragraph(); r=p.add_run(hr+'  '); r.bold=True; r.font.color.rgb=NAVY; r.font.size=Pt(9.5); r2=p.add_run(bl); r2.font.size=Pt(9.5); p.paragraph_format.space_after=Pt(1)
P(d,'En Fase 2-3 el bloque 09:00-11:00 deja de ser "deep prime de temario nuevo" y pasa a banco de preguntas + mapas conceptuales + simulacro.',c=GREY,sz=9.5)

H(d,'5 · Tabla maestra de los 40 temas (orden de rentabilidad)',12)
tb=d.add_table(rows=1,cols=5); tb.style='Light Grid Accent 1'
hc=tb.rows[0].cells
for i,h in enumerate(['Rank','Cód','Tema','Pilar','Vueltas']): hc[i].text=h
for cod,subt,ar,pr,rk,pct in sorted(TEMAS,key=lambda x:x[4]):
    c=tb.add_row().cells; c[0].text=f'#{rk}'; c[1].text=cod; c[2].text=subt; c[3].text=ar; c[4].text=f'{V[pr]} ({pr})'

H(d,'6 · Bancos / simulacros',12)
for x in ['Fuentes: QX (BanqueApp + EVA del tema + 9 simulacros + 3 oficiales) y Theomed (8 simulacros + POSTESTS por sesión).',
          'Orden: Eval diagnóstica QX (línea base) → simulacros sábado (3-4/sábado) → exámenes oficiales 2024-II/2025/2026 en la recta final.',
          'Cada simulacro: revisión brutal el mismo día (por qué la correcta + por qué se descartan los distractores). Cada fallo reinyecta su tema a +1 día.']:
    P(d,x,bullet=True,sz=10)

H(d,'7 · Top predicciones 2026-II (foco absoluto)',12)
for cod,subt,ar,pr,rk,pct in sorted(TEMAS,key=lambda x:x[4])[:10]:
    P(d,f'#{rk} · {cod} {subt} ({ar}) — ~{pct}%',bullet=True,sz=10)
d.save(os.path.join(OUT,'Plan_MAESTRO_ENCAPS_v10_25jun-20ago.docx'))

# =================== DOC 3 · SEMANAL detallado ===================
d=newdoc()
cover(d,'PLAN SEMANAL ENCAPS · SEMANA 1','Jue 25-jun → Dom 28-jun · v10 DATA-DRIVEN (detallado)')
P(d,'Inicio del plan (24-jun fue día de estructuración). Esta semana corta (jue-vie) siembra los 2 temas MÁS rentables del examen, ambos CRÍTICA (6 vueltas), 1 por día para máximo deep-work. I-5+I-6 (Bioestadística, #3) arranca el lunes 29. Sáb 27 = simulacros · Dom 28 libre.',c=GREY)
WK=[('Jueves 25-jun · D1/49','I-3','Directiva 067 RM 506-2020',
     ['Vigilancia en salud pública','Vigilancia epidemiológica','Endemias/epidemias/brotes','Demografía',
      'Historia natural del proceso salud-enfermedad','Sala situacional','Conceptos básicos de epidemiología',
      'Mediciones en epidemiología','Sistema de vigilancia epidemiológica','Gestión del riesgo en desastres']),
    ('Viernes 26-jun · D2/49','V-2','Directiva CEPLAN 001-2024 (NUEVA)',
     ['Planeamiento institucional PEI','POI y evaluación del POI','Análisis estratégico (FODA)','ROF/MOF',
      'Uso racional de medicamentos','Gestión de RRHH','Gestión logística e inventario','Buenas prácticas de almacenamiento'])]
for fecha,cod,nts,vids in WK:
    H(d,f'📅 {fecha} — {cod} {SUB[cod]}',13)
    p=d.add_paragraph()
    for tag,col in [('CRÍTICA',RED),('6 vueltas [1,3,7,28,50]',NAVY),(f'Pilar: {AR[cod]}',GREEN)]:
        r=p.add_run(f'[{tag}]  '); r.bold=True; r.font.size=Pt(9.5); r.font.color.rgb=col
    P(d,f'📋 NTS Tier-1 (obligatoria): {nts}',bold=True,c=NAVY)
    P(d,'🔴 NÚCLEO DEEP PRIME (09:00-11:00) — Videoclases QxMedic (matriz):',bold=True)
    for v in vids: P(d,v,bullet=True,sz=9.5)
    P(d,f'🎥 2ª opción: Videoclases Drive (DR LOPEZ/GALENO) del pilar {AR[cod]}. 🎯 3ª opción: clase grabada Theomed (ver minuto exacto en Joseph MD).',sz=9.5,c=GREY)
    P(d,'📄 Fichas MINSA del tema + 🎯 Compendio DR LOPEZ + 📝 POSTEST Theomed: links directos por tema en Joseph MD (pestaña ENCAPS · cada ítem con su hora).',sz=9.5,c=GREY)
    P(d,'🎯 Banco/retrieval: 20Q warm-up + 30Q consolidación + BanqueApp QX. Cada fallo → resetea a 1-3 días.',sz=9.5)
H(d,'📅 Sábado 27-jun · D3/49 — SIMULACROS',13,RED)
P(d,'Simulacros reales cronometrados + revisión brutal el mismo día:',bold=True)
for s in ['QX · Simulacro Virtual N°01','QX · Simulacro Virtual N°02','QX · Simulacro Virtual N°03']: P(d,s,bullet=True)
H(d,'📅 Domingo 28-jun · LIBRE',13,GREEN)
P(d,'Sin estudio intenso. El sueño y el modo difuso consolidan; el espaciado necesita olvido parcial. Máx 10 min re-leyendo el mapa de la semana.')
H(d,'Resumen de la semana',12,GOLD)
P(d,'2 temas CRÍTICA (los #1 y #2 del examen) sembrados con deep-work + 1 sábado de simulacros + domingo libre. La próxima semana (lun 29-jun→) acelera: I-5+I-6, II-3, III-5+II-6, I-4+II-9 y el resto del temario, 2-3 temas/día, hasta cubrir TODO al vie 17-jul.')
d.save(os.path.join(OUT,'Plan_Semanal_ENCAPS_v10_25-28jun_DATADRIVEN.docx'))

print('OK 3 Words generados en', OUT)
for f in os.listdir(OUT): print('  -', f)
