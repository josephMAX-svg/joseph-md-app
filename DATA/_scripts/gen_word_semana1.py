# -*- coding: utf-8 -*-
# Genera el Word semanal ENCAPS (Mié 24-jun → Dom 28-jun 2026), élite y detallado,
# data-driven (pronóstico 2026-2 + ciencia del aprendizaje). python3 gen_word_semana1.py
import json, re, os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
NAVY = RGBColor(0x1A, 0x2B, 0x4A); GOLD = RGBColor(0xB8, 0x86, 0x0B)
RED = RGBColor(0xC0, 0x39, 0x2B); GREEN = RGBColor(0x1E, 0x7E, 0x34); GREY = RGBColor(0x55, 0x55, 0x55)

# ---- datos reales de los días (de Supabase) ----
DIAS = [
  {"d":1,"fecha":"Miércoles 24-jun","cod":"I-3","tema":"Vigilancia Epidemiológica / Brotes","area":"Salud Pública",
   "prio":"CRÍTICA","vueltas":6,"pct":"~10.5% (RANK #1 del examen)","nts":"Directiva 067 RM 506-2020 (Vigilancia EPI)",
   "videos":["Vigilancia en salud pública","Vigilancia epidemiológica","Endemias, epidemias y brotes","Demografía",
             "Historia natural del proceso salud-enfermedad","Análisis situacional de salud (sala situacional)",
             "Conceptos básicos de epidemiología","Mediciones en epidemiología","Demografía en salud",
             "Brotes, epidemias, pandemias y endemias","Sistema de vigilancia epidemiológica",
             "Gestión del riesgo en emergencias y desastres"]},
  {"d":2,"fecha":"Jueves 25-jun","cod":"V-2","tema":"Planeamiento institucional PEI / POI / FODA","area":"Gestión",
   "prio":"CRÍTICA","vueltas":6,"pct":"~11% (RANK #2 del examen)","nts":"Directiva CEPLAN 001-2024 (NUEVA)",
   "videos":["Planeamiento institucional - PEI","Planeamiento institucional - POI y evaluación del POI",
             "Análisis estratégico institucional (FODA)","Documentos técnicos normativos de gestión (ROF/MOF)",
             "Proceso administrativo - FODA","Uso racional de medicamentos","Gestión de la historia clínica",
             "Gestión de RRHH","Gestión logística, inventario y stock de medicamentos",
             "Buenas prácticas de almacenamiento (cadena de frío)","Clima y cultura organizacional"]},
  {"d":3,"fecha":"Viernes 26-jun","cod":"I-5+I-6","tema":"Determinantes sociales + Bioestadística","area":"Salud Pública",
   "prio":"CRÍTICA","vueltas":6,"pct":"~7% (RANK #3 del examen)","nts":"DSS OMS 2008 + Bioestadística básica",
   "videos":["Causalidad y riesgo","Pruebas diagnósticas (sens/esp/VPP/VPN)",
             "Determinantes sociales - ambientales, biogenéticos y comerciales","Bioestadística básica"]},
]
SABADO = {"fecha":"Sábado 27-jun","sims":["QX · Simulacro Virtual N°01","QX · Simulacro Virtual N°02","QX · Simulacro Virtual N°03"]}

# ---- enriquecimiento desde encapsFuentes.ts (fichas + theomed sesión por tema) ----
def parse_fuentes():
    txt = open(os.path.join(ROOT,'src/lib/encapsFuentes.ts'),encoding='utf-8').read()
    fichas, sesion = {}, {}
    m = re.search(r'ENCAPS_FICHAS_POR_TEMA[^=]*=\s*(\{.*?\n\});', txt, re.S)
    if m:
        try: fichas = json.loads(m.group(1))
        except Exception: pass
    m = re.search(r'ENCAPS_THEOMED_TEMA_SESION[^=]*=\s*(\{.*?\n\});', txt, re.S)
    if m:
        try: sesion = json.loads(m.group(1))
        except Exception: pass
    return fichas, sesion
FICHAS, SESION = parse_fuentes()

doc = Document()
st = doc.styles['Normal']; st.font.name='Calibri'; st.font.size=Pt(10.5)

def H(txt, size=15, color=NAVY, after=6, before=10):
    p=doc.add_paragraph(); r=p.add_run(txt); r.bold=True; r.font.size=Pt(size); r.font.color.rgb=color
    p.paragraph_format.space_after=Pt(after); p.paragraph_format.space_before=Pt(before); return p
def P(txt, bold=False, color=None, size=10.5, bullet=False):
    p=doc.add_paragraph(style='List Bullet' if bullet else None); r=p.add_run(txt)
    r.bold=bold; r.font.size=Pt(size)
    if color: r.font.color.rgb=color
    p.paragraph_format.space_after=Pt(2); return p

# ===== PORTADA =====
t=doc.add_paragraph(); t.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=t.add_run('PLAN SEMANAL ENCAPS · SEMANA 1'); r.bold=True; r.font.size=Pt(22); r.font.color.rgb=NAVY
s=doc.add_paragraph(); s.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=s.add_run('Miércoles 24-jun → Domingo 28-jun 2026  ·  v10 DATA-DRIVEN'); r.bold=True; r.font.size=Pt(13); r.font.color.rgb=GOLD
for line in ['Joseph Max Soto Tocas · Médico, Huancayo','🎯 META: ≥17/20 · 🏥 Plaza Huachac, Junín','📅 EXAMEN: jueves 20 de agosto 2026 (tope FIJO) · 50 días hábiles',
             'Reconstruido en base a: 6 exámenes oficiales reales (2024-II→2026-I) + backtest walk-forward + ciencia del aprendizaje (Oakley · Newport · Make It Stick · Ultralearning · repetición espaciada).']:
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; rr=p.add_run(line); rr.font.size=Pt(10); rr.font.color.rgb=GREY

# ===== 0. METODOLOGÍA =====
H('0 · Por qué v10 y cómo está construido', 14)
P('Este plan ya NO es heurística. Las prioridades y vueltas salen del análisis REAL de los exámenes oficiales ENCAPS 2024-II, 2025-I, 2025-II y 2026-I (≈600 preguntas clasificadas por área/tema). Con backtest walk-forward (predije cada examen con los previos y medí el error, MAE ≈4.5pp) se construyó el pronóstico 2026-II.')
P('Front-load: los temas se ordenaron por rentabilidad. Los 3 más rentables (Vigilancia, PEI/POI, Bioestadística) ocupan los días 1-3 de esta semana para que arranquen su cadena de repaso espaciado desde HOY y alcancen 6 vueltas antes del examen.')

H('1 · Rentabilidad real proyectada 2026-II (por área)', 13)
tbl=doc.add_table(rows=1, cols=3); tbl.style='Light Grid Accent 1'
h=tbl.rows[0].cells; h[0].text='Área'; h[1].text='% examen 2026-II'; h[2].text='Vueltas / foco'
for a,pct,foco in [('I — Salud Pública','29%','MÁXIMO (Vigilancia, Bioestadística)'),
                   ('II — Cuidado Integral','28%','ALTO (Vacunas, ITS, Materna, Crónicas)'),
                   ('V — Gestión','21%','ALTO (PEI/POI, Categorización, Referencia)'),
                   ('III — Ética/Intercultural','16%','MEDIO'),
                   ('IV — Investigación','6%','BAJO (solo diseños IV-1)')]:
    c=tbl.add_row().cells; c[0].text=a; c[1].text=pct; c[2].text=foco
P('')
P('CI + SP = ~57% real (NO 70%). Investigación se desplomó a 3-6% → mínimo esfuerzo. Confianza: ranking por área fiable; % por tema son bandas ±2-4pp (n=4 exámenes).', color=GREY, size=9.5)

H('2 · Sistema de vueltas (repetición espaciada diferenciada)', 13)
for txt in ['CRÍTICA = 6 vueltas · intervalos [1, 3, 7, 28, 50] días desde el día-foco.',
            'ALTA = 5 vueltas [1, 7, 28, 50] · MEDIA = 4 [3, 28, 50] · BAJA = 3 [7, 50].',
            'Cada vuelta es RETRIEVAL (responder de memoria en el banco ANTES de ver la solución), NO relectura — el testing effect rinde 2-3× más (Make It Stick).',
            'Los intervalos se COMPRIMEN automáticamente para que toda vuelta caiga antes del 20-ago.',
            'Regla viva: cualquier tema fallado en banco/simulacro RESETEA a intervalo 1-3 días y sube a alto-yield personal.']:
    P(txt, bullet=True)

H('3 · Ritual diario (de lunes a viernes · sincronizado con Google Calendar)', 13)
for hora,bloque in [('04:15–04:45','🟢 Anclaje Post-Sueño: mapa mental del tema de ayer + validación + Anki'),
                    ('04:45–05:45','🟡 20 Preguntas Pregunta-por-Pregunta (warm-up adaptativo del tema de hoy)'),
                    ('07:15–08:15','🍅 Repaso Espaciado Multi-Temporal (D-1 / D-3 / D-7 · mapa + Anki)'),
                    ('08:15–09:00','🟡 PRE-TEST del tema del día (10 preguntas ciegas + free recall)'),
                    ('09:00–11:00','🔴 NÚCLEO DEEP PRIME (Modo A1): videos QX + compendio + 1 NTS + 8-12 APEX. Bloque de máxima energía.'),
                    ('11:00–12:00','🟡 30 Preguntas Consolidación (temas vistos · pregunta-por-pregunta + APEX)'),
                    ('17:15–18:00','🟢 Anclaje Vespertino: mapa + validación + Anki'),
                    ('18:00–18:45','🔥 Evaluación Diaria Acumulativa (modo examen + POSTEST Theomed)')]:
    p=doc.add_paragraph(); r=p.add_run(hora+'  '); r.bold=True; r.font.color.rgb=NAVY; r.font.size=Pt(9.5)
    r2=p.add_run(bloque); r2.font.size=Pt(9.5); p.paragraph_format.space_after=Pt(1)
P('Newport: máx ~4h de trabajo profundo real/día en bloques de 90 min, misma hora y lugar, teléfono fuera. Domingos LIBRES = consolidación por sueño/modo difuso (parte del método, no tiempo perdido).', color=GREY, size=9.5)

# ===== DÍA POR DÍA =====
doc.add_page_break()
H('SEMANA 1 · DÍA POR DÍA', 16, GOLD, before=0)
for D in DIAS:
    H(f"📅 {D['fecha']} · Día {D['d']}/50 — {D['cod']} {D['tema']}", 13, NAVY)
    p=doc.add_paragraph()
    for tag,col in [(f"{D['prio']}", RED if D['prio']=='CRÍTICA' else GOLD),(f"{D['vueltas']} vueltas", NAVY),(f"Rentabilidad {D['pct']}", GREEN),(f"Área: {D['area']}", GREY)]:
        rr=p.add_run(f"[{tag}]  "); rr.bold=True; rr.font.size=Pt(9.5); rr.font.color.rgb=col
    P(f"📋 NTS Tier-1 (norma obligatoria): {D['nts']}", bold=True, color=NAVY)
    P('🎬 NÚCLEO DEEP PRIME — Videoclases QxMedic (la matriz · 09:00-11:00):', bold=True)
    for v in D['videos']:
        P(v, bullet=True, size=9.5)
    # video Drive + Theomed
    P(f"🎥 2ª opción de video (Google Drive · DR LOPEZ/GALENO): Videoclases del área {D['area']}.", size=9.5, color=GREY)
    ses = SESION.get(D['cod'])
    if isinstance(ses, dict) and ses.get('sesion'):
        P(f"🎯 3ª opción — Clase grabada Theomed: este tema en Sesión {ses.get('sesion')} ({ses.get('area','')}) ≈ diapositiva {ses.get('slide')}/{ses.get('nSlides')} (~{ses.get('pct')}% del video).", size=9.5, color=GREY)
    # fichas MINSA
    fs = FICHAS.get(D['cod'], [])
    if fs:
        P(f"📄 Fichas técnicas MINSA QxMedic ({len(fs)}, ⏱ ~10 min c/u):", bold=True)
        for f in fs[:12]:
            P(f.get('titulo','') if isinstance(f,dict) else str(f), bullet=True, size=9.5)
    else:
        P("📄 Fichas MINSA del tema + 🎯 Compendio DR LOPEZ del área: ver pestaña 📚 Material en Joseph MD (link directo por tema).", size=9.5, color=GREY)
    P('🎯 Banco / práctica (retrieval): 20 preguntas warm-up + 30 consolidación + Banco QX (BanqueApp) + POSTEST Theomed del área. Cada fallo → resetea el subtema a 1-3 días.', size=9.5)
    P(f"🔁 Esta es la 1ª vuelta de {D['cod']} (CRÍTICA). Sus próximas vueltas (repaso espaciado) caen en D+1, +3, +7, +28 y +50 — todas antes del examen.", size=9.5, color=GREEN)

# Sábado
H(f"📅 {SABADO['fecha']} · Día 4/50 — SIMULACROS + revisión brutal", 13, RED)
P('Sábado de exámenes (mañana). Simulacros reales cronometrados, tamaño-examen:', bold=True)
for s in SABADO['sims']:
    P(s, bullet=True)
P('Protocolo: simulacro completo → revisión BRUTAL el mismo día (por qué la correcta lo es y por qué descartaste cada distractor). Cada fallo reinyecta su subtema al repaso de +1 día. Registra el % por área para calibrar el yield personal.', size=9.5)

# Domingo
H('📅 Domingo 28-jun · Día LIBRE (consolidación)', 13, GREEN)
P('Sin estudio intenso. El sueño y el modo difuso consolidan la memoria y el espaciado necesita olvido parcial para funcionar (Oakley/Walker). Como mucho: 10 min re-leyendo el mapa de la semana. Descanso real = parte del método.')

# Cierre
doc.add_page_break()
H('Resumen de la semana', 14, GOLD)
P('Esta semana siembra los 3 temas MÁS rentables del examen (Vigilancia EPI, PEI/POI, Bioestadística), todos CRÍTICA, iniciando su cadena de 6 vueltas desde el día 1. + 1 sábado de 3 simulacros + domingo libre.')
P('Material 100% enlazado en Joseph MD (app): cada tema con sus videos QX, video Drive de respaldo, clase grabada Theomed (con minuto), fichas MINSA, compendio DR LOPEZ, NTS y banco — cada ítem con su hora.')
P('Próxima semana (29-jun→): continúa el front-load con II-3 Vacunación, III-5 Intercultural, I-4 Defs de caso y el resto de ALTA. La app lo trae automáticamente.', color=GREY)

out_dir = r'D:\agente_estudio\ENCAPS\ENCAPS\FORMAS DE ESTUDIO\ACTUAL'
os.makedirs(out_dir, exist_ok=True)
out = os.path.join(out_dir, 'Plan_Semanal_ENCAPS_v10_24-28jun_DATADRIVEN.docx')
doc.save(out)
print('OK Word guardado:', out)
print('parse: fichas keys', len(FICHAS), '· sesion keys', len(SESION))
