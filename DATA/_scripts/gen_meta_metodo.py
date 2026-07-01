# -*- coding: utf-8 -*-
# Genera META + NIVEL + MÉTODO DIARIO ENCAPS (Word en ACTUAL + MD en DATA/ENCAPS).
# "¿Qué nivel hay que tener para >17/20?" — data-driven + escalera de hitos + prioridad + Anki.
import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = r'D:\agente_estudio\ENCAPS\ENCAPS\FORMAS DE ESTUDIO\ACTUAL'
os.makedirs(OUT, exist_ok=True)
NAVY=RGBColor(0x1A,0x2B,0x4A); GOLD=RGBColor(0xB8,0x86,0x0B); RED=RGBColor(0xC0,0x39,0x2B); GREEN=RGBColor(0x1E,0x7E,0x34); GREY=RGBColor(0x55,0x55,0x55)
doc=Document(); st=doc.styles['Normal']; st.font.name='Calibri'; st.font.size=Pt(10.5)
def H(t,sz=14,c=NAVY,af=6,bf=10):
    p=doc.add_paragraph(); r=p.add_run(t); r.bold=True; r.font.size=Pt(sz); r.font.color.rgb=c
    p.paragraph_format.space_after=Pt(af); p.paragraph_format.space_before=Pt(bf)
def P(t,bold=False,c=None,sz=10.5,bullet=False):
    p=doc.add_paragraph(style='List Bullet' if bullet else None); r=p.add_run(t); r.bold=bold; r.font.size=Pt(sz)
    if c: r.font.color.rgb=c
    p.paragraph_format.space_after=Pt(2)
md=[]
def MH(t,l=2): md.append(('#'*l)+' '+t); md.append('')
def MP(t): md.append(t); md.append('')

t=doc.add_paragraph(); t.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=t.add_run('META · NIVEL · MÉTODO DIARIO — ENCAPS 2026-II'); r.bold=True; r.font.size=Pt(19); r.font.color.rgb=NAVY
s=doc.add_paragraph(); s.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=s.add_run('¿Qué nivel hay que tener para sacar MÁS de 17/20? · plan 2-jul → examen jue 20-ago'); r.bold=True; r.font.size=Pt(12); r.font.color.rgb=GOLD
MH('META · NIVEL · MÉTODO DIARIO — ENCAPS 2026-II',1); MP('¿Qué nivel para >17/20? · plan 2-jul → examen jue 20-ago · partiendo de CERO.')

H('1 · LA META, EN NÚMEROS',13)
for x in ['>17/20 = **≥85% de aciertos** (≥85 de ~100 preguntas del examen SERUMS).',
          'Partimos de CERO hoy. 43 días de plan (2-jul→20-ago). Es alcanzable con ejecución diaria disciplinada.',
          'La nota NO se mide por sensación: se mide por **% CIEGO** (responder de memoria ANTES de ver la clave). Es la única señal real (Oakley: la ilusión de competencia mata).']:
    P(x,bullet=True)
    md.append('- '+x.replace('**',''))
md.append('')

H('2 · EL NIVEL QUE HAY QUE TENER, POR ÁREA (para que la suma ponderada dé >85%)',13)
P('Cada área pesa distinto en el examen (backtest de 6 exámenes reales). Para llegar a 85.6% global, este es el % que debes dominar en cada pilar:',c=GREY,sz=9.5)
tb=doc.add_table(rows=1,cols=4); tb.style='Light Grid Accent 1'
for i,h in enumerate(['Pilar','% del examen','Nivel objetivo (% ciego)','Foco']): tb.rows[0].cells[i].text=h
filas=[('I · Salud Pública','29%','≥ 90%','Vigilancia, Bioestadística, Defs de caso'),
       ('II · Cuidado Integral','28%','≥ 88%','Vacunas, ITS, Materna, Crónicas'),
       ('V · Gestión','21%','≥ 85%','PEI/POI, Categorización, Referencia'),
       ('III · Ética/Intercultural','16%','≥ 80%','Intercultural, Derechos/Deberes'),
       ('IV · Investigación','6%','≥ 70%','Solo diseños (IV-1)')]
for f in filas:
    c=tb.add_row().cells
    for i,v in enumerate(f): c[i].text=v
P('')
P('Suma ponderada: 0.29×90 + 0.28×88 + 0.21×85 + 0.16×80 + 0.06×70 = 85.6% → >17/20. ✔ Si un área baja, otra debe subir para mantener el 85%.',bold=True,c=GREEN,sz=10)
MH('El nivel por área (para >85% global)')
MP('| Pilar | % examen | Nivel objetivo (% ciego) |\n|---|---|---|\n| I Salud Pública | 29% | ≥90% |\n| II Cuidado Integral | 28% | ≥88% |\n| V Gestión | 21% | ≥85% |\n| III Ética/Intercultural | 16% | ≥80% |\n| IV Investigación | 6% | ≥70% |\n\nSuma ponderada = 85.6% → >17/20.')

H('3 · ESCALERA DE HITOS (cómo saber si vas en nivel para el 20-ago)',13)
P('Mide tu % CIEGO en los simulacros/Eval y compáralo con el mínimo de cada fase. Si estás por debajo → reajusta (más banco de ese tema, resetea sus vueltas).',c=GREY,sz=9.5)
tb=doc.add_table(rows=1,cols=4); tb.style='Light List Accent 1'
for i,h in enumerate(['Fase','Fechas','Qué haces','% mínimo (ciego)']): tb.rows[0].cells[i].text=h
hitos=[('FASE 1 · COBERTURA','2-jul → 24-jul','1ª vuelta de TODO el temario (17 temas). Cada tema cierra con su Eval.','≥60% en la Eval del tema (recién visto)'),
       ('FASE 2 · VUELTAS+PREGUNTAS','27-jul → 13-ago','Repaso espaciado + banco + simulacros. 65-70% del tiempo al alto yield.','≥75% en simulacros completos · ≥85% en I/II/V'),
       ('FASE 3 · RECTA FINAL','14-ago → 19-ago','SOLO simulacros tamaño-examen + preguntas. Cero temario nuevo.','≥85-90% consistente · CRÍTICA ≥90%'),
       ('EXAMEN','jue 20-ago','—','META: >17/20 (≥85%)')]
for h in hitos:
    c=tb.add_row().cells
    for i,v in enumerate(h): c[i].text=v
P('')
P('Regla de oro (Make It Stick): cada fallo en banco/simulacro RESETEA ese subtema a repaso de 1-3 días y sube a tu "alto-yield personal". Los fallos son tu mapa.',bold=True,c=RED,sz=10)
MH('Escalera de hitos (% ciego mínimo por fase)')
MP('- Fase 1 (2-24 jul): ≥60% en la Eval del tema recién visto.\n- Fase 2 (27-jul→13-ago): ≥75% simulacros · ≥85% en I/II/V.\n- Fase 3 (14-19 ago): ≥85-90% consistente · CRÍTICA ≥90%.\n- Examen 20-ago: >17/20 (≥85%).')

H('4 · MÉTODO DIARIO — qué leer/hacer cada día (8 bloques, ya en Joseph MD + Calendar)',13)
for hr,bl in [('04:15','🟢 Anclaje Post-Sueño: mapa mental del tema de ayer + Anki (repaso FSRS)'),
              ('04:45','🟡 20 Preguntas warm-up del tema de hoy (pregunta-por-pregunta, CIEGAS)'),
              ('07:15','🍅 Repaso Espaciado D-1/D-3/D-7 (mapas + Anki de temas previos)'),
              ('08:15','🟡 PRE-TEST: 10 preguntas ciegas del tema + free recall (qué recuerdas sin mirar)'),
              ('09:00-11:00','🔴 NÚCLEO DEEP PRIME: videos QX del tema + 1 NTS + compendio DR LOPEZ + fichas MINSA (8-12 APEX). El bloque de máxima energía.'),
              ('11:00','🟡 30 Preguntas Consolidación (banco QX + POSTEST Theomed) + PRÁCTICA EXTRA (Kahoot/Banqueo)'),
              ('17:15','🟢 Anclaje Vespertino: mapa + validación + Anki'),
              ('18:00','🔥 Evaluación Diaria: modo examen (Eval del tema · registra tu % ciego en la app)')]:
    p=doc.add_paragraph(); r=p.add_run(hr+'  '); r.bold=True; r.font.color.rgb=NAVY; r.font.size=Pt(9.5)
    r2=p.add_run(bl); r2.font.size=Pt(9.5); p.paragraph_format.space_after=Pt(1)
P('El detalle EXACTO de cada día (qué video, qué ficha, qué banco, con su hora y link) está en Joseph MD → Estudio → ENCAPS (pestaña ⚡ Hoy). Esto es el ritual; la app pone el contenido del día.',c=GREY,sz=9.5)
MH('Método diario (8 bloques · en la app + Calendar)')
MP('04:15 Anclaje+Anki · 04:45 20Q warm-up ciegas · 07:15 Repaso D-1/3/7 · 08:15 PRE-TEST 10Q+free recall · **09:00-11:00 Núcleo Deep Prime** (videos QX+NTS+compendio+fichas) · 11:00 30Q consolidación+práctica extra · 17:15 Anclaje vespertino · 18:00 Evaluación diaria (registra % ciego). El contenido exacto del día = app Joseph MD → ENCAPS → ⚡ Hoy.')

H('5 · PRIORIDAD ENCAPS · qué hacer si te atrasas',13)
for x in ['ENCAPS es la PRIORIDAD ABSOLUTA hasta el 20-ago. Todo lo demás se subordina.',
          'Si un día NO terminas el tema ENCAPS: **reclama los bloques de MIR / USMLE / Research / AURUM / Synapse del Google Calendar** (12:30 Synapse, 13:30 Research/Derma, 14:15 AURUM, 15:15 MIR, 16:15 USMLE) y úsalos para cerrar ENCAPS. Esos segmentos pueden esperar; el examen del 20-ago no.',
          'Regla nueva (tú lo pediste): NO habrá más reprogramaciones. Solo si me avisas "me atrasé en X porque Y" reestructuro — con el porqué, para no repetir.']:
    P(x,bullet=True)
    md.append('- '+x.replace('**',''))
md.append('')

H('6 · ANKI — cómo se maneja',13)
for x in ['Anki es el motor de retención (bloques Anclaje 04:15 y 17:15 + Repaso 07:15).',
          '1 idea por tarjeta (Palmerton/Yousmle): nada de tarjetas con párrafos. Pregunta atómica → respuesta atómica.',
          'FSRS activado (programación óptima). Los temas CRÍTICA (I-3, V-2, I-5+I-6, II-3) reciben MÁS tarjetas.',
          'Cada fallo en banco/simulacro → crea/refuerza su tarjeta Anki ese mismo día.',
          'Meta diaria Anki: 0 tarjetas atrasadas al dormir (el "Again/Good" es Anki estándar, no método propio).']:
    P(x,bullet=True); md.append('- '+x.replace('**',''))
md.append('')

H('7 · VERIFICACIÓN DE FUENTES (todo dentro, nada fuera · en vivo)',12)
P('QxMedic: 5 áreas · 105 videoclases · Evaluaciones (Diagnóstica/EVA/Simulacros). Theomed: 3 cursos (73 regular con áreas+asincrónicas 1-8+en vivo+repasos+banqueos · 37 simulacros · 89 kahoots). Google Drive: 5 proveedores (DR LOPEZ con Compendio/Normativas/Simulacros/Kahoot). TODO mapeado en la app: 40 temas · 161 videos QX · 105 fichas MINSA · 35 evaluaciones + práctica extra. Cero temas/subtemas fuera.',sz=9.5)
MH('Fuentes verificadas en vivo'); MP('QX 5 áreas/105 videoclases · Theomed 3 cursos (73/37/89) · Drive 5 proveedores. 40 temas · 161 videos · 105 fichas · 35 evaluaciones + práctica extra. Nada fuera.')

out_docx=os.path.join(OUT,'META_NIVEL_METODO_ENCAPS_v10.docx'); doc.save(out_docx)
out_md=os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))),'ENCAPS','META_NIVEL_METODO_ENCAPS.md')
open(out_md,'w',encoding='utf-8').write('\n'.join(md))
print('OK Word:',out_docx); print('OK MD:',out_md)
